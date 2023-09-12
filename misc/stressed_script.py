from tkinter import *
from tkinter import filedialog
from tkinter import messagebox
from pathlib import Path
from docx import Document
import os

MOST_WANTED = "ѝ"
times_found = 0


def open_file():
    global times_found
    chosen_file = filedialog.askopenfilename(parent=main_window, title="Please select a file:",
                                             filetypes=(("Word Document", "*.docx"), ("Plain Text", "*.txt"), ))

    file_path = Path(chosen_file)

    if file_path.suffix == ".docx":
        process_docx(file_path)
    elif file_path.suffix == ".txt":
        process_txt(file_path)

    if times_found == 0:
        messagebox.showinfo("Статус", "Не открих нищо!")
    elif times_found == 1:
        messagebox.showinfo("Статус",
                            "Открито е едно ударено и!\nВ същата папка ще откриеш нов файл с маркираното място.")
    elif times_found > 1:
        messagebox.showinfo("Статус",
                            f"Открити са {times_found} ударени и-та!\nВ същата папка ще откриеш нов файл с маркираните места.")

    times_found = 0


def process_docx(file_path):
    global times_found
    doc = Document(file_path)

    tables = doc.tables
    paragraphs = doc.paragraphs

    for table in tables:
        for cell in table._cells:
            if MOST_WANTED in cell.text:
                times_found += len([word for word in cell.text.split() if MOST_WANTED in word])
                cell.text = cell.text.replace("ѝ", "**ѝ**")

    for paragraph in paragraphs:
        if MOST_WANTED in paragraph.text:
            paragraph.text = paragraph.text.replace("ѝ", "**ѝ**")
            times_found += len([word for word in paragraph.text.split() if MOST_WANTED in word])

    file_folder = file_path.parent
    file_name = file_path.stem
    doc.save(f"{file_folder}/{file_name}_checked.docx")


def process_txt(file_path):
    global times_found
    file_folder = file_path.parent
    file_name = file_path.stem
    output_file = f"{file_folder}/{file_name}_checked.txt"

    if os.path.isfile(output_file):
        os.remove(output_file)

    with open(file_path, "r", encoding="utf8") as reader, open(output_file, "w", encoding="utf8") as writer:
        for line in reader.readlines():
            writer.write(line.replace("ѝ", f"**{MOST_WANTED}**"))
            times_found += len([word for word in line.split() if MOST_WANTED in word])


main_window = Tk()
main_window.geometry()
main_window.title("Удрѝ" + u"\u2122")
main_window.resizable(width="False", height="False")

instructions = """Провери текстови файл за скрити ударени и-та.\n
Скриптът отбелязва позициите им и запазва резултата в нов файл.\n
Настоящата версия работи с файлове в .docx и .txt формат."""

instructions_label = Label(main_window, text=instructions, justify="left", pady=10)
open_button = Button(main_window, text="Провери файл", width=27, command=open_file)
closing_button = Button(main_window, text="Затвори прозореца", width=27, command=main_window.destroy)

instructions_label.pack()
open_button.pack(side="left")
closing_button.pack(side="left")

main_window.mainloop()